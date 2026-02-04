import os
import argparse
import sys
import re
from intelligence import AnalysisAgent
from ppt_generator import PPTGenerator
from visual_engine import VisualEngine
from data_loader import UniversalLoader

# 1. ENV VAR CHECK
GEMINI_KEY = os.getenv("GEMINI_API_KEY") or "YOUR_GEMINI_KEY"
PEXELS_KEY = os.getenv("PEXELS_API_KEY") or "YOUR_PEXELS_KEY"

# Graceful docx import
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

if "YOUR_" in GEMINI_KEY:
    print("❌ ERROR: Please set GEMINI_API_KEY environment variable.")
    sys.exit(1)

def generate_citation_doc(strategy_data, chunks, output_path):
    if not HAS_DOCX:
        with open(output_path.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
            f.write(f"CITATION REPORT - {strategy_data.get('code_name')}\n")
            for c in strategy_data.get('citations', []):
                f.write(f"CLAIM: {c.get('claim')}\nSOURCE: {c.get('source_display')}\nID: {c.get('id')}\n\n")
        return

    doc = Document()
    doc.add_heading(f"Citation Document - {strategy_data.get('code_name')}", 0)
    
    chunk_map = {c['id']: c for c in chunks}
    
    for cite in strategy_data.get('citations', []):
        claim = cite.get('claim', 'Claim')
        cid = cite.get('id')
        
        doc.add_heading(claim, level=2)
        if cid in chunk_map:
            c = chunk_map[cid]
            doc.add_paragraph(f"Source: {c['source']}")
            doc.add_paragraph(f"Location: {c['location']}")
            doc.add_paragraph(f"Excerpt: \"{c['text'][:300]}...\"")
        else:
            doc.add_paragraph(f"Source: {cite.get('source_display', 'Unknown')}")
            doc.add_paragraph("Note: Direct chunk reference not found in data vault.")
            
    doc.save(output_path)
    print(f"✅ Citation Doc saved: {output_path}")

def assess_data_quality(chunks):
    quality = {
        "private": len([c for c in chunks if 'private' in c['type']]),
        "web": len([c for c in chunks if 'public' in c['type']]),
        "financial": len([c for c in chunks if 'financial' in c['type']]),
        "total": len(chunks)
    }
    print(f"📊 DATA QUALITY: {quality['total']} chunks (Pvt: {quality['private']}, Web: {quality['web']}, Fin: {quality['financial']})")
    return quality

def clean_company_name(filename):
    """Extracts 'Ind Swift' from 'Ind Swift-OnePager.md'"""
    base = os.path.splitext(os.path.basename(filename))[0]
    # Remove common suffixes to get the REAL name for checking leaks
    base = re.sub(r'[-_ ]?(OnePager|Pitch|Deck|Teaser|Report|Analysis)', '', base, flags=re.IGNORECASE)
    return base.strip()

def process_company(file_path, loader, agent, visual, builder):
    c_name = clean_company_name(file_path)
    print(f"\n🚀 Processing: {c_name} (File: {os.path.basename(file_path)})")
    
    # A. Ingest Private Data
    chunks = loader.load_data(file_path)
    
    # B. EXTRACT PUBLIC URL
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
            url_match = re.search(r'##\s*Website.*?(https?://[^\s<>\)\"]+)', content, re.IGNORECASE | re.DOTALL)
            if url_match:
                target_url = url_match.group(1)
                print(f"🌍 Found Website: {target_url} -> Scraping...")
                web_chunks = loader.load_data(target_url)
                chunks.extend(web_chunks)
    except Exception as e:
        print(f"⚠️ Warning: URL extraction error: {e}")

    if not chunks:
        print("❌ No data found.")
        return {"success": False, "company": c_name, "cost": 0}

    assess_data_quality(chunks)

    # C. Analyze
    data = agent.analyze_company(chunks, c_name)
    
    if not data: 
        print("❌ Agent Analysis Failed (Check logs above).")
        return {"success": False, "company": c_name, "cost": agent.cost_tracker.session_cost}

    # D. Visuals
    imgs = []
    sec = data.get('sector', 'General')
    kws = data.get('visual_keywords', ['business'])
    for i in range(3):
        kw = kws[i] if i < len(kws) else 'office'
        u = visual.fetch_image(kw, sec, slide_index=i+1)
        if u:
            fn = f"temp_{c_name}_{i}.jpg"
            visual.download_image(u, fn)
            imgs.append(fn)
        else: imgs.append(None)

    # E. Outputs
    out_ppt = f"Output_{c_name}.pptx"
    out_doc = f"Citations_{c_name}.docx" if HAS_DOCX else f"Citations_{c_name}.txt"
    
    builder.generate_ppt(data, imgs, out_ppt)
    generate_citation_doc(data, chunks, out_doc)
    
    # F. Cleanup Temps
    for img in imgs:
        if img and os.path.exists(img):
            try: os.remove(img)
            except: pass
    
    return {"success": True, "company": c_name, "cost": agent.cost_tracker.session_cost}

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--file", help="Single file")
    parser.add_argument("--folder", help="Batch folder")
    args = parser.parse_args()

    # 1. Initialize
    loader = UniversalLoader()
    agent = AnalysisAgent(GEMINI_KEY)
    visual = VisualEngine(PEXELS_KEY)
    builder = PPTGenerator()
    
    # 2. Pre-flight Check
    if not agent.test_api_connection():
        sys.exit(1)

    # 3. Process
    results = []
    if args.folder:
        for f in os.listdir(args.folder):
            if f.endswith(('.md', '.pdf', '.docx', '.xlsx')):
                agent.cost_tracker.session_cost = 0 
                res = process_company(os.path.join(args.folder, f), loader, agent, visual, builder)
                results.append(res)
    elif args.file:
        res = process_company(args.file, loader, agent, visual, builder)
        results.append(res)
    else:
        if os.path.exists("Centum-OnePager.md"):
            res = process_company("Centum-OnePager.md", loader, agent, visual, builder)
            results.append(res)

    print("\n" + "="*50)
    print("BATCH PROCESSING SUMMARY")
    total_cost = 0
    for r in results:
        status = "✅" if r['success'] else "❌"
        cost = r['cost']
        total_cost += cost
        print(f"{status} {r['company']}: ₹{cost:.2f}")
    print(f"TOTAL RUN COST: ₹{total_cost:.2f}")

if __name__ == "__main__":
    main()